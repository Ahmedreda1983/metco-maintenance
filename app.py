from flask import Flask, render_template, request, jsonify, send_from_directory, send_file
import pandas as pd
from datetime import datetime
import os
import sqlite3
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import shutil
import re
import io

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

app = Flask(__name__)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ASSET_FILE_PATH = os.path.join(BASE_DIR, "Asset List.xlsx")
PM_FILE_PATH = os.path.join(BASE_DIR, "PM List.xlsx")

UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
IMAGES_FOLDER = os.path.join(UPLOAD_FOLDER, 'Images')
os.makedirs(IMAGES_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def clean_filename(name):
    return re.sub(r'[\\/:*?"<>|#]', '_', name.strip().replace(' ', '_'))

pm_search_columns = [
    "Work Order", "PM", "Job Plan", "Parent WO", "Description", "Location", "Asset",
    "MMS #", "QR CODE", "Route", "Work Type", "Workshop", "Target Start", "Target Finish",
    "METCO COMMENT"
]

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def load_all_sheets(file_path, header_row):
    if not os.path.exists(file_path):
        return {}
    xls = pd.ExcelFile(file_path)
    return {sheet: pd.read_excel(file_path, sheet_name=sheet, header=header_row) for sheet in xls.sheet_names}

def init_db():
    conn = sqlite3.connect('data.db')
    c = conn.cursor()
    c.execute('''
    CREATE TABLE IF NOT EXISTS maintenance_records (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        type TEXT,
        sheet_name TEXT,
        row_index INTEGER,
        data TEXT,
        timestamp TEXT,
        before_images TEXT,
        after_images TEXT,
        report_images TEXT,
        cm_images TEXT,
        notes_text TEXT,
        notes_images TEXT
    )''')
    c.execute('''
    CREATE TABLE IF NOT EXISTS zip_files (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        filename TEXT,
        content BLOB
    )''')
    conn.commit()
    conn.close()

init_db()

def upload_to_drive(file_path, filename):
    try:
        SCOPES = ['https://www.googleapis.com/auth/drive.file']
        SERVICE_ACCOUNT_FILE = os.path.join(BASE_DIR, 'credentials.json')
        credentials = service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
        service = build('drive', 'v3', credentials=credentials)
        folder_id = '1SolVjxUU0iZ7YRgmgt424_JMtkPU6CUG'
        media = MediaFileUpload(file_path, mimetype='application/zip')
        file_metadata = {'name': filename, 'parents': [folder_id]}
        file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
        print(f"File uploaded to Google Drive with ID: {file.get('id')}")
        return True
    except Exception as e:
        print("Error uploading to Google Drive:", e)
        return False

def save_zip_to_db(filename, filepath):
    try:
        with open(filepath, 'rb') as f:
            content = f.read()
        conn = sqlite3.connect('data.db')
        c = conn.cursor()
        c.execute("INSERT INTO zip_files (filename, content) VALUES (?, ?)", (filename, content))
        conn.commit()
        conn.close()
        print(f"Saved ZIP to database: {filename}")
    except Exception as e:
        print("Error saving ZIP to database:", e)

# === Main saving handler for both PM and Asset ===
def handle_save(req, entry_type):
    try:
        sheet_name = req.form.get('sheet_name')
        row_index = int(req.form.get('row_index'))
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        data = {key[6:]: req.form.get(key) for key in req.form if key.startswith('field_')}
        notes_text = req.form.get('notes_text', '')

        def save_files(files_list):
            saved_files = []
            for file in files_list:
                if file and allowed_file(file.filename):
                    filename = secure_filename(f"{datetime.now().strftime('%Y%m%d%H%M%S%f')}_{file.filename}")
                    save_path = os.path.join(IMAGES_FOLDER, filename)
                    file.save(save_path)
                    saved_files.append(filename)
            return saved_files

        before_images = save_files(req.files.getlist('Before Maintenance')) if entry_type == 'PM' else []
        after_images = save_files(req.files.getlist('After Maintenance')) if entry_type == 'PM' else []
        report_images = save_files(req.files.getlist('Maintenance Report')) if entry_type == 'PM' else []
        cm_images = save_files(req.files.getlist('CM Images')) if entry_type in ['CM', 'Asset'] else []
        spare_parts_images = save_files(req.files.getlist('Spare Parts Images')) if entry_type == 'Asset' else []
        notes_images = save_files(req.files.getlist('notes_images'))

        # Save to DB
        conn = sqlite3.connect('data.db')
        c = conn.cursor()
        c.execute('''
            INSERT INTO maintenance_records (type, sheet_name, row_index, data, timestamp, before_images, after_images, report_images, cm_images, notes_text, notes_images)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            entry_type, sheet_name, row_index, json.dumps(data, ensure_ascii=False), timestamp,
            ','.join(before_images), ','.join(after_images), ','.join(report_images),
            ','.join(cm_images), notes_text, ','.join(notes_images)
        ))
        conn.commit()
        conn.close()

        # Export Excel and Word
        data['ملاحظات'] = notes_text
        df = pd.DataFrame([data])
        df = df[[col for col in pm_search_columns if col in df.columns] + ['ملاحظات']]

        wo = clean_filename(data.get('Work Order') or data.get('Asset') or 'UnknownWO')
        location = clean_filename(data.get('Location') or 'UnknownLoc')
        desc = clean_filename(data.get('Description') or 'NoDesc')[:30]
        file_timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{wo}_{file_timestamp}_{location}_{desc}"

        output_folder = os.path.join(UPLOAD_FOLDER, base_filename)
        os.makedirs(output_folder, exist_ok=True)

        excel_file = os.path.join(output_folder, f"{base_filename}.xlsx")
        df.to_excel(excel_file, sheet_name='PM_Records', index=False)

        word_path = os.path.join(output_folder, f"{base_filename}.docx")
        doc = Document()
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)
        doc.add_heading(f"Work Order: {data.get('Work Order', data.get('Asset', ''))}", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=0, cols=2)
        for key, val in data.items():
            if val:
                row = table.add_row().cells
                row[0].text = str(key)
                row[1].text = str(val)
        doc.add_paragraph()

        def add_images_table(images_list, label):
            if not images_list: return
            doc.add_paragraph(label, style='Heading 2')
            table = doc.add_table(rows=0, cols=3)
            for i in range(0, len(images_list), 3):
                row_cells = table.add_row().cells
                for j, img in enumerate(images_list[i:i+3]):
                    img_path = os.path.join(IMAGES_FOLDER, img)
                    if os.path.exists(img_path):
                        row_cells[j].paragraphs[0].add_run().add_picture(img_path, width=Inches(2))

        add_images_table(report_images, 'Report')
        add_images_table(before_images, 'Before')
        add_images_table(after_images, 'After')
        add_images_table(cm_images, 'CM Images')
        add_images_table(spare_parts_images, 'Spare Parts')
        if notes_text:
            doc.add_paragraph('ملاحظات:', style='Heading 2')
            doc.add_paragraph(notes_text)
        add_images_table(notes_images, 'Notes')

        doc.save(word_path)

        for img in before_images + after_images + report_images + cm_images + notes_images + spare_parts_images:
            shutil.move(os.path.join(IMAGES_FOLDER, img), os.path.join(output_folder, img))

        zip_filename = f"{base_filename}.zip"
        zip_path = os.path.join(UPLOAD_FOLDER, zip_filename)
        shutil.make_archive(zip_path.replace('.zip', ''), 'zip', output_folder)
        upload_to_drive(zip_path, zip_filename)
        save_zip_to_db(zip_filename, zip_path)

        return jsonify({"message": "Saved", "zip_file": zip_filename})
    except Exception as e:
        print("ERROR in handle_save:", e)
        return jsonify({"error": str(e)}), 500
@app.route('/')
def index():
    search_query = request.args.get('search', '').strip().lower()
    conn = sqlite3.connect('data.db')
    c = conn.cursor()
    if search_query:
        c.execute("SELECT id, filename FROM zip_files WHERE LOWER(filename) LIKE ? ORDER BY id DESC", (f"%{search_query}%",))
    else:
        c.execute("SELECT id, filename FROM zip_files ORDER BY id DESC")
    zip_files = c.fetchall()
    conn.close()
    return render_template('index.html', zip_files=zip_files)

@app.route('/download_zip/<int:zip_id>')
def download_zip(zip_id):
    try:
        conn = sqlite3.connect('data.db')
        c = conn.cursor()
        c.execute("SELECT filename, content FROM zip_files WHERE id=?", (zip_id,))
        row = c.fetchone()
        conn.close()
        if row:
            filename, content = row
            return send_file(io.BytesIO(content), as_attachment=True, download_name=filename, mimetype='application/zip')
        else:
            return "ZIP not found", 404
    except Exception as e:
        return f"Error: {e}", 500

@app.route('/search', methods=['POST'])
def search():
    try:
        query = request.json.get('query', '').strip().lower()
        if not query:
            return jsonify({"assets": [], "pm": []})

        asset_sheets = load_all_sheets(ASSET_FILE_PATH, header_row=1)
        pm_sheets = load_all_sheets(PM_FILE_PATH, header_row=0)

        asset_results, pm_results = [], []
        for sheet_name, df in asset_sheets.items():
            df = df.fillna('')
            for idx, row in df.iterrows():
                if any(query in str(value).lower() for value in row):
                    asset_results.append({"SheetName": sheet_name, "RowIndex": int(idx), "data": row.to_dict()})

        for sheet_name, df in pm_sheets.items():
            df = df.fillna('')
            for idx, row in df.iterrows():
                if any(query in str(row[col]).lower() for col in pm_search_columns if col in df.columns):
                    pm_results.append({"SheetName": sheet_name, "RowIndex": int(idx), "data": row.to_dict()})

        return jsonify({"assets": asset_results, "pm": pm_results})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/get_record')
def get_record():
    try:
        data_type = request.args.get("type")
        sheet_name = request.args.get("sheet")
        row_index = int(request.args.get("row"))
        sheets = load_all_sheets(PM_FILE_PATH if data_type == "PM" else ASSET_FILE_PATH, 0 if data_type == "PM" else 1)
        df = sheets.get(sheet_name)
        if df is None:
            return jsonify({"error": "Sheet not found"}), 404
        row = df.iloc[row_index].fillna('').to_dict()
        return jsonify({"data": row, "sheet": sheet_name, "row": row_index, "type": data_type})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/save_edit', methods=['POST'])
def save_edit():
    return handle_save(request, entry_type='PM')

@app.route('/save_asset_edit', methods=['POST'])
def save_asset_edit():
    return handle_save(request, entry_type='Asset')

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)
@app.route('/search_zip', methods=['POST'])
def search_zip():
    try:
        query = request.json.get('query', '').strip().lower()
        conn = sqlite3.connect('data.db')
        cursor = conn.cursor()
        cursor.execute("SELECT filename FROM zip_files WHERE LOWER(filename) LIKE ?", (f"%{query}%",))
        files = [row[0] for row in cursor.fetchall()]
        conn.close()
        return jsonify(files)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/edit')
def edit_page():
    return render_template('edit_page.html')
@app.route('/edit_asset')
def edit_asset_page():
    sheet = request.args.get('sheet')
    row = request.args.get('row')
    return render_template('edit_asset.html', sheet_name=sheet, row_index=row)

@app.route('/download_zip_by_name/<path:filename>')
def download_zip_by_name(filename):
    try:
        conn = sqlite3.connect('data.db')
        c = conn.cursor()
        c.execute("SELECT content FROM zip_files WHERE filename = ?", (filename,))
        row = c.fetchone()
        conn.close()
        if row:
            content = row[0]
            return send_file(io.BytesIO(content), as_attachment=True, download_name=filename, mimetype='application/zip')
        else:
            return "ZIP file not found", 404
    except Exception as e:
        return f"Error: {e}", 500
@app.route('/download_zip/<path:filename>')
def download_zip_filename(filename):
    try:
        conn = sqlite3.connect('data.db')
        c = conn.cursor()
        c.execute("SELECT content FROM zip_files WHERE filename = ?", (filename,))
        row = c.fetchone()
        conn.close()
        if row:
            content = row[0]
            return send_file(io.BytesIO(content), as_attachment=True, download_name=filename, mimetype='application/zip')
        else:
            return "ZIP file not found", 404
    except Exception as e:
        return f"Error: {e}", 500


# لازم ترجع للسطر ده لتشغيل السيرفر:
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
